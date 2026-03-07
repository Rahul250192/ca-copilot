"""
Agreement PDF Generation Service
Generates professional PDF documents from agreement templates.
"""
import io
import re
from typing import Dict, List, Any
from datetime import datetime


def _replace_placeholders(text: str, values: Dict[str, Any]) -> str:
    """Replace [PLACEHOLDER] tokens in clause text with actual values."""
    result = text
    # Map common placeholder patterns to field names
    mapping = {
        "[FIRM NAME]": values.get("firm_name", "[FIRM NAME]"),
        "[DESCRIPTION]": values.get("description", values.get("business_nature", "[DESCRIPTION]")),
        "[DATE]": values.get("agreement_date", datetime.now().strftime("%d/%m/%Y")),
        "[PARTY A]": values.get("party_a_name", "[PARTY A]"),
        "[PARTY B]": values.get("party_b_name", "[PARTY B]"),
        "[PARTY A NAME]": values.get("party_a_name", "[PARTY A NAME]"),
        "[PARTY B NAME]": values.get("party_b_name", "[PARTY B NAME]"),
        "[LANDLORD NAME]": values.get("party_a_name", "[LANDLORD NAME]"),
        "[TENANT NAME]": values.get("party_b_name", "[TENANT NAME]"),
        "[LANDLORD]": values.get("party_a_name", "[LANDLORD]"),
        "[TENANT]": values.get("party_b_name", "[TENANT]"),
        "[LESSOR]": values.get("party_a_name", "[LESSOR]"),
        "[LESSEE]": values.get("party_b_name", "[LESSEE]"),
        "[LICENSOR]": values.get("party_a_name", "[LICENSOR]"),
        "[LICENSEE]": values.get("party_b_name", "[LICENSEE]"),
        "[SELLER]": values.get("party_a_name", "[SELLER]"),
        "[BUYER]": values.get("party_b_name", "[BUYER]"),
        "[LENDER]": values.get("party_a_name", "[LENDER]"),
        "[BORROWER]": values.get("party_b_name", "[BORROWER]"),
        "[INVESTOR]": values.get("party_a_name", "[INVESTOR]"),
        "[COMPANY]": values.get("company_name", values.get("firm_name", "[COMPANY]")),
        "[COMPANY NAME]": values.get("company_name", values.get("firm_name", "[COMPANY NAME]")),
        "[SERVICE PROVIDER]": values.get("party_a_name", "[SERVICE PROVIDER]"),
        "[CLIENT]": values.get("party_b_name", "[CLIENT]"),
        "[CONSULTANT NAME]": values.get("party_a_name", "[CONSULTANT NAME]"),
        "[DISCLOSING PARTY]": values.get("party_a_name", "[DISCLOSING PARTY]"),
        "[RECEIVING PARTY]": values.get("party_b_name", "[RECEIVING PARTY]"),
        "[EMPLOYEE NAME]": values.get("party_b_name", values.get("employee_name", "[EMPLOYEE NAME]")),
        "[INTERN NAME]": values.get("party_b_name", "[INTERN NAME]"),
        "[VENDOR]": values.get("party_a_name", "[VENDOR]"),
        "[SUPPLIER]": values.get("party_a_name", "[SUPPLIER]"),
        "[DISTRIBUTOR]": values.get("party_b_name", "[DISTRIBUTOR]"),
        "[AGENT]": values.get("party_b_name", "[AGENT]"),
        "[PRINCIPAL]": values.get("party_a_name", "[PRINCIPAL]"),
        "[FULL ADDRESS]": values.get("address", "[FULL ADDRESS]"),
        "[ADDRESS]": values.get("address", "[ADDRESS]"),
        "[CITY]": values.get("city", "[CITY]"),
        "[AMOUNT]": values.get("amount", "[AMOUNT]"),
        "[AREA]": values.get("area", "[AREA]"),
        "[START DATE]": values.get("start_date", values.get("agreement_date", "[START DATE]")),
        "[END DATE]": values.get("end_date", "[END DATE]"),
        "[DESIGNATION]": values.get("designation", "[DESIGNATION]"),
        "[DEPARTMENT]": values.get("department", "[DEPARTMENT]"),
        "[LOCATION]": values.get("location", values.get("city", "[LOCATION]")),
        "[MANAGER NAME/TITLE]": values.get("manager_name", "[MANAGER NAME/TITLE]"),
        "[REPORTING MANAGER]": values.get("manager_name", "[REPORTING MANAGER]"),
        "[MENTOR NAME]": values.get("mentor_name", "[MENTOR NAME]"),
        "[LLP NAME]": values.get("firm_name", "[LLP NAME]"),
        "[TERRITORY]": values.get("territory", "[TERRITORY]"),
        "[PRODUCTS]": values.get("products", "[PRODUCTS]"),
        "[PRODUCTS/SERVICES]": values.get("products", values.get("services_desc", "[PRODUCTS/SERVICES]")),
        "[PROFESSIONAL/FIRM]": values.get("party_a_name", "[PROFESSIONAL/FIRM]"),
        "[NUMBER]": values.get("registration_number", "[NUMBER]"),
        "[SUBJECT]": values.get("subject", "[SUBJECT]"),
        "[X]": values.get("notice_period", values.get("term_months", "[X]")),
        "[Y]": values.get("profit_share_b", "[Y]"),
    }

    for placeholder, value in mapping.items():
        result = result.replace(placeholder, str(value))

    return result


def generate_agreement_html(
    agreement_name: str,
    agreement_desc: str,
    field_values: Dict[str, Any],
    clauses: List[Dict[str, Any]],
) -> str:
    """Generate an HTML representation of the agreement for preview and PDF conversion."""
    today = datetime.now().strftime("%B %d, %Y")
    date_val = field_values.get("agreement_date", today)

    # Build clause HTML
    clause_blocks = []
    for i, clause in enumerate(clauses):
        if not clause.get("is_default", True):
            continue  # skip unchecked clauses
        content = _replace_placeholders(clause.get("content", ""), field_values)
        # Convert newlines to <br> for proper rendering
        content_html = content.replace("\n", "<br>")
        clause_blocks.append(f"""
        <div class="clause">
            <h3>{i + 1}. {clause.get('title', '')}</h3>
            <p>{content_html}</p>
        </div>
        """)

    clauses_html = "\n".join(clause_blocks)

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap');
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: 'Outfit', 'Georgia', serif;
    color: #1a1a2e; line-height: 1.7;
    padding: 48px 56px; max-width: 800px; margin: 0 auto;
    background: white;
}}
.header {{
    text-align: center; margin-bottom: 36px;
    border-bottom: 2px solid #1e3a8a; padding-bottom: 24px;
}}
.header h1 {{
    font-size: 22px; font-weight: 700; color: #1e3a8a;
    text-transform: uppercase; letter-spacing: 1.5px; margin-bottom: 6px;
}}
.header .date {{ font-size: 13px; color: #64748b; }}
.clause {{
    margin-bottom: 20px; page-break-inside: avoid;
}}
.clause h3 {{
    font-size: 14px; font-weight: 700; color: #1e3a8a;
    margin-bottom: 6px; text-transform: uppercase; letter-spacing: 0.5px;
}}
.clause p {{
    font-size: 13px; color: #334155; text-align: justify;
}}
.signature-block {{
    margin-top: 60px; display: flex; justify-content: space-between;
    page-break-inside: avoid;
}}
.sig-party {{
    width: 45%; text-align: center;
}}
.sig-line {{
    border-top: 1px solid #1e293b; margin-top: 60px; padding-top: 8px;
    font-size: 12px; color: #64748b;
}}
.sig-name {{ font-weight: 600; font-size: 13px; color: #1e293b; }}
.footer {{
    margin-top: 40px; text-align: center; font-size: 11px;
    color: #94a3b8; border-top: 1px solid #e2e8f0; padding-top: 16px;
}}
</style>
</head>
<body>
<div class="header">
    <h1>{agreement_name}</h1>
    <div class="date">Date: {date_val}</div>
</div>

{clauses_html}

<div class="signature-block">
    <div class="sig-party">
        <div class="sig-line">
            <div class="sig-name">{field_values.get('party_a_name', 'First Party')}</div>
            Authorized Signatory
        </div>
    </div>
    <div class="sig-party">
        <div class="sig-line">
            <div class="sig-name">{field_values.get('party_b_name', 'Second Party')}</div>
            Authorized Signatory
        </div>
    </div>
</div>

<div class="footer">
    This document was generated using ComplianceAI Expert Platform.<br>
    Stamp duty and registration as applicable under Indian law.
</div>
</body>
</html>"""

    return html


# ═══════════════════════════════════════════════════════
# DOCX TEMPLATE ENGINE
# ═══════════════════════════════════════════════════════

# Paragraph-index → ordered list of field names for each blank in that paragraph.
# This maps the Rent_Lease_Agreement.docx template blanks to field values.
RENT_LEASE_BLANK_MAP = {
    2:  ["agreement_day", "agreement_month", "agreement_year", "city"],
    6:  ["landlord_name", "landlord_age", "landlord_parent", "landlord_address"],
    10: ["tenant_name", "tenant_age", "tenant_parent", "tenant_address"],
    18: ["flat_no", "floor", "building"],
    19: ["street"],
    20: ["premise_city", "premise_state", "premise_pin"],
    21: ["area_sqft"],
    22: ["survey_no"],
    29: ["purpose"],
    34: ["start_date", "term_months", "end_date"],
    38: ["monthly_rent", "rent_words", "rent_due_day"],
    39: ["bank_name", "account_no", "ifsc", "upi_id"],
    40: ["rent_increment_pct"],
    41: ["late_fee_pct"],
    43: ["security_deposit", "deposit_words"],
    44: ["deposit_refund_days"],
    47: ["purpose"],
    57: ["maintenance_charges"],
    63: ["notice_period_months"],
    64: ["default_months"],
    67: ["lockin_months"],
    77: ["police_verify_days"],
    82: ["jurisdiction_city"],
    89: ["num_originals"],
}

# Paragraph-index → ordered list of field names for Retainer_Agreement_India.docx
RETAINER_BLANK_MAP = {
    2:  ["agreement_day", "agreement_month", "agreement_year", "city"],
    6:  ["client_name", "client_address"],
    10: ["retainer_name", "retainer_designation", "retainer_address"],
    23: ["service_a"],
    24: ["service_b"],
    25: ["service_c"],
    31: ["start_date", "term_months", "end_date"],
    34: ["retainer_fee", "retainer_fee_words", "fee_due_day"],
    35: ["bank_name", "account_no", "ifsc", "upi_id"],
    36: ["fee_revision_pct"],
    43: ["confidentiality_years"],
    54: ["non_solicit_months"],
    55: ["non_solicit_client_months"],
    80: ["termination_notice_months"],
    82: ["default_months"],
    87: ["arbitration_city"],
    88: ["jurisdiction_city"],
    92: ["num_originals"],
}

# Paragraph-index → ordered list of field names for LLP_Agreement_Template.docx
LLP_BLANK_MAP = {
    15: ["llp_name"],
    18: ["execution_date", "effective_date", "execution_place"],
    21: ["partner_1_name", "partner_1_father", "partner_1_address"],
    23: ["partner_2_name", "partner_2_father", "partner_2_address"],
    27: ["partner_3_name", "partner_3_father", "partner_3_address"],
    33: ["llp_name_repeat", "llp_id"],
    36: ["llp_business_desc"],
    112: ["llp_name_repeat_2"],
    116: ["llp_registered_office"],
    555: ["signature_year", "signature_place"],
}
# There are also blanks in tables 0, 2, 5. They will be handled generically in the table processing.

MOU_BLANK_MAP = {
    10: ["party_a_name", "party_a_address"],
    13: ["party_a_business", "party_b_business"],
    15: ["collaboration_field"],
    22: ["mou_purpose"],
    26: ["effective_date", "valid_months"],
    29: ["party_a_responsibility_1"],
    30: ["party_a_responsibility_2"],
    31: ["party_a_responsibility_3"],
    33: ["party_b_responsibility_1"],
    34: ["party_b_responsibility_2"],
    35: ["party_b_responsibility_3"],
    41: ["confidentiality_years"],
    48: ["jurisdiction_city"],
    51: ["num_originals"],
}

SERVICE_BLANK_MAP = {
    2: ["execution_day", "execution_month", "execution_year", "execution_city"],
    6: ["client_name", "client_address"],
    10: ["provider_name", "provider_address"],
    21: ["service_1"],
    22: ["service_2"],
    23: ["service_3"],
    27: ["start_date", "end_date"],
    29: ["fee_amount_numbers", "fee_amount_words"],
    30: ["advance_pct"],
    32: ["final_pct"],
    33: ["invoice_days", "interest_pct"],
    35: ["out_of_pocket_details"],
    38: ["review_days"],
    39: ["resubmit_days"],
    41: ["confidentiality_survive_years"],
    53: ["termination_notice_days"],
    54: ["cure_period_days"],
    59: ["arbitration_city", "courts_city"],
    63: ["num_originals"],
}

OFFER_LETTER_BLANK_MAP = {
    2: ["ref_no"],
    3: ["offer_date"],
    6: ["employee_name"],
    7: ["employee_father_name"],
    8: ["employee_address"],
    10: ["position_title"],
    12: ["employee_name_salutation"],
    15: ["company_name"],
    19: ["designation"],
    20: ["department"],
    21: ["reporting_to"],
    22: ["posting_location"],
    24: ["joining_date"],
    26: ["ctc_numbers", "ctc_words"],
    28: ["salary_credit_day"],
    30: ["probation_months"],
    33: ["hours_per_day", "days_per_week"],
    43: ["notice_period_months"],
    47: ["non_compete_months"],
    49: ["governing_state"],
    51: ["acceptance_days"],
    62: ["hr_name"],
    63: ["hr_designation"],
    64: ["company_name_footer"],
    69: ["employee_name_accepting"],
    72: ["employee_signature", "date_signed"],
    74: ["employee_name_footer", "pan_number"],
}

COMMISSION_BLANK_MAP = {
    4: ["party_a_name", "party_a_address", "party_b_name", "party_b_address"],
    6: ["agent_name", "company_name", "territory", "client_name", "product"],
    8: ["termination_months"],
    10: ["commission_rate"],
    12: ["effective_date", "start_date"],
    14: ["agent_territory"],
    25: ["arbitration_country"],
    33: ["company_name_signature"],
    37: ["date_signed"],
}

# Template path → blank_map routing
TEMPLATE_BLANK_MAPS = {
    "property_and_lease/Rent_Lease_Agreement.docx": RENT_LEASE_BLANK_MAP,
    "Rent_Lease_Agreement.docx": RENT_LEASE_BLANK_MAP,  # legacy path
    "service_and_professionals/Retainer_Agreement_India.docx": RETAINER_BLANK_MAP,
    "business_and_partnership/LLP_Agreement_Template.docx": LLP_BLANK_MAP,
    "business_and_partnership/MOU_India.docx": MOU_BLANK_MAP,
    "service_and_professionals/Service_Agreement_India.docx": SERVICE_BLANK_MAP,
    "employment_and_hr/Offer_Letter_India.docx": OFFER_LETTER_BLANK_MAP,
    "vendor_and_procurement/Commission_Agreement_Template.docx": COMMISSION_BLANK_MAP,
}


def get_blank_map_for_template(template_path: str = None) -> dict:
    """Return the correct blank_map for a given template path."""
    if template_path and template_path in TEMPLATE_BLANK_MAPS:
        return TEMPLATE_BLANK_MAPS[template_path]
    return RENT_LEASE_BLANK_MAP  # fallback


def _replace_blanks_in_paragraph(paragraph, field_values: list):
    """
    Replace _{3,} runs in a paragraph with actual values.
    field_values is a list of values in order of blank appearance.
    """
    import re
    # Reconstruct the full text to know blank positions
    full_text = paragraph.text
    blank_pattern = re.compile(r'_{3,}')
    blanks_found = list(blank_pattern.finditer(full_text))

    if not blanks_found or not field_values:
        return

    # Build a mapping of character offset → replacement value
    replacements = {}
    for i, match in enumerate(blanks_found):
        if i < len(field_values):
            val = str(field_values[i]) if field_values[i] else "________________"
            replacements[match.start()] = (match.end(), val)

    if not replacements:
        return

    # Now walk through runs and replace blanks
    # We need to track our offset in the full text
    offset = 0
    for run in paragraph.runs:
        run_text = run.text
        run_start = offset
        run_end = offset + len(run_text)

        new_text = ""
        pos = 0  # position within run_text

        for blank_start, (blank_end, val) in sorted(replacements.items()):
            # Check if this blank overlaps with this run
            if blank_start >= run_end or blank_end <= run_start:
                continue

            # Calculate positions within the run
            local_start = max(0, blank_start - run_start)
            local_end = min(len(run_text), blank_end - run_start)

            # Add text before the blank
            if local_start > pos:
                new_text += run_text[pos:local_start]

            # Only add replacement value if this is the first run containing this blank
            if blank_start >= run_start:
                new_text += val

            pos = local_end

        # Add remaining text in run
        if pos < len(run_text):
            new_text += run_text[pos:]

        run.text = new_text
        offset = run_end


def generate_docx_from_template(
    template_bytes: bytes,
    field_values: Dict[str, Any],
    blank_map: Dict[int, list] = None,
    template_path: str = None,
) -> bytes:
    """
    Fill a .docx template by replacing underline blanks with field values.

    Args:
        template_bytes: Raw bytes of the .docx template
        field_values: Dict of field_name → value
        blank_map: Optional paragraph-index → [field_names] mapping.
                   Defaults to RENT_LEASE_BLANK_MAP.

    Returns:
        Filled .docx as bytes
    """
    import docx

    if blank_map is None:
        blank_map = get_blank_map_for_template(template_path)

    doc = docx.Document(io.BytesIO(template_bytes))

    for para_idx, field_names in blank_map.items():
        if para_idx >= len(doc.paragraphs):
            continue

        paragraph = doc.paragraphs[para_idx]
        # Resolve field names to values
        values = []
        for fname in field_names:
            val = field_values.get(fname, "")
            if not val and fname == "purpose":
                val = field_values.get("usage_purpose", "")
            values.append(val if val else "________________")

        _replace_blanks_in_paragraph(paragraph, values)

    # Save to buffer
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


# ═══════════════════════════════════════════════════════
# DOCX → HTML PREVIEW (with named placeholder markers)
# ═══════════════════════════════════════════════════════

def _run_to_html(run) -> str:
    """Convert a single docx Run to HTML with formatting."""
    import html as html_mod
    text = html_mod.escape(run.text)
    if not text:
        return ""
    if run.bold:
        text = f"<b>{text}</b>"
    if run.italic:
        text = f"<i>{text}</i>"
    if run.underline:
        text = f"<u>{text}</u>"
    return text


def _para_to_html(paragraph, para_idx: int, blank_map: Dict[int, list]) -> str:
    """Convert a paragraph to HTML, inserting named placeholder spans for blanks."""
    import html as html_mod

    # Determine tag based on style
    style_name = ""
    if paragraph.style and paragraph.style.name:
        style_name = paragraph.style.name.lower()
    if "heading 1" in style_name:
        tag = "h1"
    elif "heading 2" in style_name:
        tag = "h2"
    elif "heading 3" in style_name:
        tag = "h3"
    else:
        tag = "p"

    # Get blank field names for this paragraph
    field_names = blank_map.get(para_idx, [])

    if not field_names:
        # No placeholders — render runs directly
        inner = "".join(_run_to_html(r) for r in paragraph.runs)
        if not inner.strip():
            return ""
        return f"<{tag}>{inner}</{tag}>"

    # Has blanks — reconstruct text and replace _{3,} with placeholder spans
    full_text = paragraph.text
    blank_pattern = re.compile(r'_{3,}')
    parts = []
    last_end = 0
    blank_idx = 0

    for match in blank_pattern.finditer(full_text):
        # Text before the blank
        before = full_text[last_end:match.start()]
        parts.append(html_mod.escape(before))

        # Placeholder span
        if blank_idx < len(field_names):
            fname = field_names[blank_idx]
            label = fname.replace("_", " ").title()
            parts.append(
                f'<span class="tpl-ph" data-field="{fname}">{label}</span>'
            )
            blank_idx += 1
        else:
            parts.append('<span class="tpl-ph">________________</span>')

        last_end = match.end()

    # Remaining text
    if last_end < len(full_text):
        parts.append(html_mod.escape(full_text[last_end:]))

    inner = "".join(parts)
    return f"<{tag}>{inner}</{tag}>"


def _table_to_html(table) -> str:
    """Convert a docx table to HTML."""
    import html as html_mod
    rows_html = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cells.append(f"<td>{html_mod.escape(cell.text)}</td>")
        rows_html.append(f"<tr>{''.join(cells)}</tr>")
    return f'<table class="tpl-table">{"".join(rows_html)}</table>'


def docx_to_preview_html(
    template_bytes: bytes,
    blank_map: Dict[int, list] = None,
    template_path: str = None,
) -> str:
    """
    Convert a .docx template to styled HTML for live preview.
    Underline blanks are replaced with <span data-field="..."> markers
    that can be swapped client-side for live editing.

    Returns full HTML document string.
    """
    import docx

    if blank_map is None:
        blank_map = get_blank_map_for_template(template_path)

    doc = docx.Document(io.BytesIO(template_bytes))

    # Build body HTML from paragraphs
    # We need to interleave paragraphs and tables in document order
    body_parts = []

    # python-docx doesn't expose document-order iteration easily,
    # so we iterate paragraphs with index and insert tables at their positions
    para_idx = 0
    for element in doc.element.body:
        tag_name = element.tag.split("}")[-1] if "}" in element.tag else element.tag
        if tag_name == "p":
            if para_idx < len(doc.paragraphs):
                html_part = _para_to_html(doc.paragraphs[para_idx], para_idx, blank_map)
                if html_part:
                    body_parts.append(html_part)
                para_idx += 1
        elif tag_name == "tbl":
            # Find corresponding table
            for table in doc.tables:
                if table._element is element:
                    body_parts.append(_table_to_html(table))
                    break

    body_html = "\n".join(body_parts)

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&display=swap');
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: 'Outfit', 'Georgia', serif;
    color: #1a1a2e; line-height: 1.8;
    padding: 40px 48px; max-width: 800px; margin: 0 auto;
    background: white; font-size: 13px;
}}
h1 {{
    font-size: 18px; font-weight: 800; color: #1e3a8a;
    text-align: center; text-transform: uppercase;
    letter-spacing: 1.5px; margin: 0 0 4px;
    border-bottom: 2px solid #1e3a8a; padding-bottom: 10px;
}}
h2 {{
    font-size: 14px; font-weight: 700; color: #1e3a8a;
    margin: 20px 0 6px; text-transform: uppercase;
    letter-spacing: 0.5px;
}}
h3 {{
    font-size: 13px; font-weight: 700; color: #334155;
    margin: 14px 0 4px;
}}
p {{
    margin: 6px 0; text-align: justify; color: #334155;
}}
b, strong {{ font-weight: 700; }}
.tpl-ph {{
    display: inline;
    background: #eff6ff;
    color: #1e40af;
    border: 1px dashed #93c5fd;
    border-radius: 4px;
    padding: 1px 6px;
    font-weight: 600;
    font-size: 12px;
    white-space: nowrap;
    transition: all 0.2s;
}}
.tpl-ph.filled {{
    background: #f0fdf4;
    color: #166534;
    border-color: #86efac;
    font-weight: 700;
}}
.tpl-table {{
    width: 100%; border-collapse: collapse;
    margin: 16px 0; font-size: 12px;
}}
.tpl-table td {{
    border: 1px solid #e2e8f0;
    padding: 6px 10px; color: #334155;
    vertical-align: top;
}}
.tpl-table tr:first-child td {{
    background: #f8fafc; font-weight: 700;
    color: #1e293b;
}}
</style>
</head>
<body>
{body_html}
</body>
</html>"""

